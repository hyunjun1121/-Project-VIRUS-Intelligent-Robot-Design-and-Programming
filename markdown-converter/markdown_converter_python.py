#!/usr/bin/env python3
# -*- coding: utf-8 -*-
r"""
Markdown Converter Tool (Python 라이브러리 버전)
마크다운 파일의 이스케이프된 볼드 텍스트(\*\*)를 제대로 된 볼드 텍스트(**)로 변환하고
PDF와 DOCX 형식으로 내보내는 도구 (Python 라이브러리 사용)
"""

import re
import os
import sys
import argparse
from pathlib import Path

def fix_escaped_bold(text):
    r"""
    모든 백슬래시(\)를 제거하여 이스케이프된 마크다운 문법을 수정
    """
    # 모든 백슬래시 제거
    text = text.replace('\\', '')
    
    return text

def convert_to_pdf_weasyprint(markdown_content, output_path):
    """
    마크다운 내용을 PDF로 변환 (weasyprint 사용)
    """
    try:
        import markdown
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
        
        # 마크다운을 HTML로 변환
        md = markdown.Markdown(extensions=['extra', 'codehilite'])
        html_content = md.convert(markdown_content)
        
        # HTML 템플릿 작성
        full_html = f"""
        <!DOCTYPE html>
        <html lang="ko">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Converted Document</title>
            <style>
                body {{
                    font-family: "Malgun Gothic", "맑은 고딕", "Nanum Gothic", "나눔고딕", sans-serif;
                    font-size: 12pt;
                    line-height: 1.6;
                    margin: 2cm;
                    color: #333;
                }}
                h1, h2, h3, h4, h5, h6 {{
                    color: #2c3e50;
                    margin-top: 1.5em;
                    margin-bottom: 0.5em;
                }}
                code {{
                    background-color: #f8f9fa;
                    padding: 0.2em 0.4em;
                    border-radius: 3px;
                    font-family: "Consolas", "Monaco", monospace;
                }}
                pre {{
                    background-color: #f8f9fa;
                    padding: 1em;
                    border-radius: 5px;
                    overflow-x: auto;
                }}
                blockquote {{
                    border-left: 4px solid #3498db;
                    padding-left: 1em;
                    margin-left: 0;
                    color: #7f8c8d;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 1em 0;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }}
                th {{
                    background-color: #f2f2f2;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # PDF 생성
        font_config = FontConfiguration()
        html_doc = HTML(string=full_html)
        html_doc.write_pdf(output_path, font_config=font_config)
        
        print(f"✅ PDF 파일이 성공적으로 생성되었습니다: {output_path}")
        return True
        
    except ImportError as e:
        print(f"❌ 필요한 라이브러리가 설치되지 않았습니다: {str(e)}")
        print("설치 명령: pip install markdown weasyprint")
        return False
    except Exception as e:
        print(f"❌ PDF 변환 중 오류 발생: {str(e)}")
        print("💡 Windows에서 PDF 생성이 어려울 수 있습니다. HTML 파일을 생성해서 브라우저에서 PDF로 저장하세요.")
        # HTML 파일 생성 시도
        try:
            import markdown
            md = markdown.Markdown(extensions=['extra', 'codehilite'])
            html_content = md.convert(markdown_content)
            
            full_html = f"""
            <!DOCTYPE html>
            <html lang="ko">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Converted Document</title>
                <style>
                    body {{
                        font-family: "Malgun Gothic", "맑은 고딕", "Nanum Gothic", "나눔고딕", sans-serif;
                        font-size: 12pt;
                        line-height: 1.6;
                        margin: 2cm;
                        color: #333;
                    }}
                    h1, h2, h3, h4, h5, h6 {{
                        color: #2c3e50;
                        margin-top: 1.5em;
                        margin-bottom: 0.5em;
                    }}
                    code {{
                        background-color: #f8f9fa;
                        padding: 0.2em 0.4em;
                        border-radius: 3px;
                        font-family: "Consolas", "Monaco", monospace;
                    }}
                    pre {{
                        background-color: #f8f9fa;
                        padding: 1em;
                        border-radius: 5px;
                        overflow-x: auto;
                    }}
                    blockquote {{
                        border-left: 4px solid #3498db;
                        padding-left: 1em;
                        margin-left: 0;
                        color: #7f8c8d;
                    }}
                    table {{
                        border-collapse: collapse;
                        width: 100%;
                        margin: 1em 0;
                    }}
                    th, td {{
                        border: 1px solid #ddd;
                        padding: 8px;
                        text-align: left;
                    }}
                    th {{
                        background-color: #f2f2f2;
                    }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            html_path = output_path.replace('.pdf', '.html')
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
            print(f"✅ HTML 파일을 생성했습니다: {html_path}")
            print("   브라우저에서 열어서 '인쇄 → PDF로 저장'을 선택하세요.")
            return True
        except:
            return False

def convert_to_docx_python(markdown_content, output_path):
    """
    마크다운 내용을 DOCX로 변환 (python-docx 사용)
    """
    try:
        import markdown
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.style import WD_STYLE_TYPE
        from bs4 import BeautifulSoup
        
        # 마크다운을 HTML로 변환
        md = markdown.Markdown(extensions=['extra'])
        html_content = md.convert(markdown_content)
        
        # BeautifulSoup으로 HTML 파싱
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 새 문서 생성
        doc = Document()
        
        # 기본 폰트 설정
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Malgun Gothic'
        font.size = Pt(11)
        
        # HTML 요소들을 DOCX로 변환
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'blockquote', 'table']):
            if element.name.startswith('h'):
                # 헤딩 처리
                level = int(element.name[1])
                heading = doc.add_heading(element.get_text(), level=level)
            elif element.name == 'p':
                # 문단 처리
                paragraph = doc.add_paragraph()
                add_formatted_text(paragraph, element)
            elif element.name in ['ul', 'ol']:
                # 리스트 처리
                for li in element.find_all('li'):
                    p = doc.add_paragraph(li.get_text(), style='List Bullet' if element.name == 'ul' else 'List Number')
            elif element.name == 'blockquote':
                # 인용문 처리
                quote = doc.add_paragraph(element.get_text())
                quote.style = 'Quote'
            elif element.name == 'table':
                # 테이블 처리 (간단한 버전)
                rows = element.find_all('tr')
                if rows:
                    cols = len(rows[0].find_all(['th', 'td']))
                    table = doc.add_table(rows=len(rows), cols=cols)
                    table.style = 'Table Grid'
                    
                    for i, row in enumerate(rows):
                        cells = row.find_all(['th', 'td'])
                        for j, cell in enumerate(cells):
                            if j < cols:
                                table.cell(i, j).text = cell.get_text()
        
        # 문서 저장
        doc.save(output_path)
        print(f"✅ DOCX 파일이 성공적으로 생성되었습니다: {output_path}")
        return True
        
    except ImportError as e:
        print(f"❌ 필요한 라이브러리가 설치되지 않았습니다: {str(e)}")
        print("설치 명령: pip install markdown python-docx beautifulsoup4")
        return False
    except Exception as e:
        print(f"❌ DOCX 변환 중 오류 발생: {str(e)}")
        return False

def add_formatted_text(paragraph, element):
    """
    HTML 요소의 서식을 DOCX 문단에 추가
    """
    for content in element.contents:
        if hasattr(content, 'name'):
            if content.name == 'strong' or content.name == 'b':
                run = paragraph.add_run(content.get_text())
                run.bold = True
            elif content.name == 'em' or content.name == 'i':
                run = paragraph.add_run(content.get_text())
                run.italic = True
            elif content.name == 'code':
                run = paragraph.add_run(content.get_text())
                run.font.name = 'Consolas'
            else:
                paragraph.add_run(content.get_text())
        else:
            paragraph.add_run(str(content))

def process_markdown_file(input_path, output_dir=None, formats=['pdf', 'docx']):
    """
    마크다운 파일을 처리하고 지정된 형식으로 변환
    """
    input_path = Path(input_path)
    
    if not input_path.exists():
        print(f"❌ 입력 파일을 찾을 수 없습니다: {input_path}")
        return False
    
    # 출력 디렉토리 설정
    if output_dir is None:
        output_dir = input_path.parent
    else:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
    
    # 마크다운 파일 읽기
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        print(f"❌ 파일 읽기 오류: {str(e)}")
        return False
    
    # 이스케이프된 볼드 텍스트 수정
    print("🔧 이스케이프된 마크다운 문법을 수정하는 중...")
    fixed_content = fix_escaped_bold(content)
    
    # 수정된 마크다운 파일 저장
    fixed_md_path = output_dir / f"{input_path.stem}_fixed.md"
    try:
        with open(fixed_md_path, 'w', encoding='utf-8') as f:
            f.write(fixed_content)
        print(f"✅ 수정된 마크다운 파일 저장: {fixed_md_path}")
    except Exception as e:
        print(f"⚠️ 수정된 마크다운 파일 저장 실패: {str(e)}")
    
    success = True
    
    # PDF 변환
    if 'pdf' in formats:
        pdf_path = output_dir / f"{input_path.stem}.pdf"
        if not convert_to_pdf_weasyprint(fixed_content, str(pdf_path)):
            success = False
    
    # DOCX 변환
    if 'docx' in formats:
        docx_path = output_dir / f"{input_path.stem}.docx"
        if not convert_to_docx_python(fixed_content, str(docx_path)):
            success = False
    
    return success

def main():
    parser = argparse.ArgumentParser(
        description='마크다운 파일의 이스케이프된 볼드 텍스트를 수정하고 PDF/DOCX로 변환 (Python 라이브러리 버전)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
사용 예시:
  python markdown_converter_python.py input.md
  python markdown_converter_python.py input.md -o output_folder
  python markdown_converter_python.py input.md -f pdf docx
  python markdown_converter_python.py input.md -o output_folder -f pdf

필요한 라이브러리 설치:
  pip install markdown weasyprint python-docx beautifulsoup4
        """
    )
    
    parser.add_argument('input', help='입력 마크다운 파일 경로')
    parser.add_argument('-o', '--output', help='출력 디렉토리 (기본값: 입력 파일과 같은 디렉토리)')
    parser.add_argument('-f', '--formats', nargs='+', choices=['pdf', 'docx'], 
                       default=['pdf', 'docx'], help='출력 형식 (기본값: pdf docx)')
    
    args = parser.parse_args()
    
    print("🚀 마크다운 변환기 시작 (Python 라이브러리 버전)...")
    print(f"📄 입력 파일: {args.input}")
    print(f"📁 출력 디렉토리: {args.output or '입력 파일과 같은 디렉토리'}")
    print(f"📋 출력 형식: {', '.join(args.formats)}")
    print("-" * 50)
    
    success = process_markdown_file(args.input, args.output, args.formats)
    
    if success:
        print("\n✅ 변환이 완료되었습니다!")
    else:
        print("\n❌ 변환 중 일부 오류가 발생했습니다.")
        sys.exit(1)

if __name__ == "__main__":
    main() 